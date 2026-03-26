"""
Document API routes
"""
import os
import shutil
import logging
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import List, Optional
from pydantic import BaseModel

from config import settings
from app.utils.database import get_db
from app.models.database import Document, Intent
from app.models.schemas import DocumentResponse, DocumentListResponse, DocumentUploadResponse

logger = logging.getLogger(__name__)
router = APIRouter()


@router.get("", response_model=DocumentListResponse)
async def list_documents(
    skip: int = 0,
    limit: int = 100,
    intent_id: Optional[int] = None,
    status: Optional[str] = None,
    search: Optional[str] = None,
    file_type: Optional[str] = None,
    db: AsyncSession = Depends(get_db)
):
    """List all documents with filters"""
    query = select(Document)

    # Apply filters
    if intent_id is not None:
        query = query.where(Document.intent_id == intent_id)
    if status is not None:
        query = query.where(Document.status == status)
    if search:
        query = query.where(Document.name.contains(search))
    if file_type is not None:
        query = query.where(Document.file_type == file_type)

    # Get total count
    count_result = await db.execute(
        select(func.count()).select_from(query.subquery())
    )
    total = count_result.scalar() or 0

    # Get documents
    query = query.order_by(Document.created_at.desc()).offset(skip).limit(limit)
    result = await db.execute(query)
    documents = result.scalars().all()

    # Get intent names
    items = []
    for doc in documents:
        intent_name = None
        if doc.intent_id:
            intent_result = await db.execute(
                select(Intent.name).where(Intent.id == doc.intent_id)
            )
            intent_name = intent_result.scalar_one_or_none()

        items.append(DocumentResponse(
            id=doc.id,
            name=doc.name,
            file_path=doc.file_path,
            file_size=doc.file_size,
            file_type=doc.file_type,
            intent_id=doc.intent_id,
            intent_name=intent_name,
            status=doc.status,
            error_message=doc.error_message,
            created_at=doc.created_at,
            updated_at=doc.updated_at
        ))

    return DocumentListResponse(items=items, total=total)


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Get document by ID"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    intent_name = None
    if doc.intent_id:
        intent_result = await db.execute(
            select(Intent.name).where(Intent.id == doc.intent_id)
        )
        intent_name = intent_result.scalar_one_or_none()

    return DocumentResponse(
        id=doc.id,
        name=doc.name,
        file_path=doc.file_path,
        file_size=doc.file_size,
        file_type=doc.file_type,
        intent_id=doc.intent_id,
        intent_name=intent_name,
        status=doc.status,
        error_message=doc.error_message,
        created_at=doc.created_at,
        updated_at=doc.updated_at
    )


@router.get("/{document_id}/download")
async def download_document(
    document_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Download document file"""
    from fastapi.responses import FileResponse
    
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    media_type = "application/pdf" if doc.file_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return FileResponse(
        path=doc.file_path,
        filename=doc.name,
        media_type=media_type
    )


class DocumentContentResponse(BaseModel):
    document_id: int
    document_name: str
    content: str
    preview: str
    word_count: int


@router.get("/{document_id}/content", response_model=DocumentContentResponse)
async def get_document_content(
    document_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Get document text content"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # First, try to use content from database (already parsed during upload)
    if doc.content:
        logger.info(f"Using stored content for document {document_id}")
        content = doc.content
    else:
        # Fallback: extract from file if no stored content
        if not os.path.exists(doc.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        content = ""
        
        try:
            if doc.file_type == "docx":
                import docx2txt
                content = docx2txt.process(doc.file_path)
            elif doc.file_type == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(doc.file_path)
                for page in reader.pages:
                    content += page.extract_text() or ""
        except Exception as e:
            logger.error(f"Error extracting content: {e}")
            content = f"[Unable to extract content: {str(e)}]"
    
    # Create preview (first 500 chars)
    preview = content[:500] + "..." if len(content) > 500 else content
    word_count = len(content.split())
    
    return DocumentContentResponse(
        document_id=doc.id,
        document_name=doc.name,
        content=content,
        preview=preview,
        word_count=word_count
    )


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    intent_id: Optional[int] = Form(None),
    db: AsyncSession = Depends(get_db)
):
    """Upload a document (PDF or DOCX)"""
    allowed_types = {"pdf": "pdf", "docx": "docx"}
    file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""

    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {list(allowed_types.keys())}"
        )

    if intent_id:
        intent_result = await db.execute(
            select(Intent).where(Intent.id == intent_id)
        )
        if not intent_result.scalar_one_or_none():
            raise HTTPException(status_code=400, detail="Invalid intent_id")

    import uuid
    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    file_path = settings.UPLOADS_DIR / unique_filename

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    file_size = file_path.stat().st_size

    document = Document(
        name=file.filename,
        file_path=str(file_path),
        file_size=file_size,
        file_type=allowed_types[file_ext],
        intent_id=intent_id,
        status="pending"
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    from app.services.tasks import process_document_async
    import asyncio
    asyncio.create_task(process_document_async(document.id))

    return DocumentUploadResponse(
        id=document.id,
        name=document.name,
        status=document.status,
        message="Document uploaded successfully. Processing will begin automatically."
    )


class BatchUploadResponse(BaseModel):
    total: int
    successful: int
    failed: int
    results: List[dict]


@router.post("/upload-batch", response_model=BatchUploadResponse)
async def upload_documents_batch(
    files: List[UploadFile] = File(...),
    intent_id: Optional[int] = Form(None),
    db: AsyncSession = Depends(get_db)
):
    """Upload multiple documents (PDF or DOCX)"""
    allowed_types = {"pdf": "pdf", "docx": "docx"}
    results = []
    successful = 0
    failed = 0

    if intent_id:
        intent_result = await db.execute(
            select(Intent).where(Intent.id == intent_id)
        )
        if not intent_result.scalar_one_or_none():
            raise HTTPException(status_code=400, detail="Invalid intent_id")

    import uuid
    from app.services.tasks import process_document_async
    import asyncio

    for file in files:
        try:
            file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""

            if file_ext not in allowed_types:
                results.append({
                    "name": file.filename,
                    "status": "failed",
                    "error": f"Invalid file type. Allowed: pdf, docx"
                })
                failed += 1
                continue

            unique_filename = f"{uuid.uuid4()}_{file.filename}"
            file_path = settings.UPLOADS_DIR / unique_filename

            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            file_size = file_path.stat().st_size

            document = Document(
                name=file.filename,
                file_path=str(file_path),
                file_size=file_size,
                file_type=allowed_types[file_ext],
                intent_id=intent_id,
                status="pending"
            )
            db.add(document)
            await db.commit()
            await db.refresh(document)

            asyncio.create_task(process_document_async(document.id))

            results.append({
                "id": document.id,
                "name": document.name,
                "status": "uploaded",
                "message": "Uploaded. Processing will begin automatically."
            })
            successful += 1

        except Exception as e:
            results.append({
                "name": file.filename,
                "status": "failed",
                "error": str(e)
            })
            failed += 1

    await db.commit()

    return BatchUploadResponse(
        total=len(files),
        successful=successful,
        failed=failed,
        results=results
    )


@router.put("/{document_id}/intent")
async def update_document_intent(
    document_id: int,
    intent_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Update document's intent association"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate intent_id
    intent_result = await db.execute(
        select(Intent).where(Intent.id == intent_id)
    )
    if not intent_result.scalar_one_or_none():
        raise HTTPException(status_code=400, detail="Invalid intent_id")

    document.intent_id = intent_id
    await db.commit()

    return {"message": "Document intent updated successfully"}


@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Delete a document"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file from disk
    if os.path.exists(document.file_path):
        os.remove(document.file_path)

    # Delete from vector store
    from app.services.document_service import delete_document_from_vector_store
    await delete_document_from_vector_store(document_id)

    await db.delete(document)
    await db.commit()

    return {"message": "Document deleted successfully"}


@router.post("/{document_id}/reparse")
async def reparse_document(
    document_id: int,
    chunk_size: Optional[int] = Form(None),
    chunk_overlap: Optional[int] = Form(None),
    rechunk: bool = Form(True),
    db: AsyncSession = Depends(get_db)
):
    """Reparse a document with optional new chunk settings"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Update chunk settings if provided
    if chunk_size is not None:
        settings.CHUNK_SIZE = chunk_size
    if chunk_overlap is not None:
        settings.CHUNK_OVERLAP = chunk_overlap

    # Reset status to pending for reprocessing
    document.status = "pending"
    document.vector_ids = []
    await db.commit()

    # Trigger reprocessing
    from app.services.tasks import process_document_async
    import asyncio
    asyncio.create_task(process_document_async(document.id, force_rechunk=rechunk))

    return {
        "message": "Document marked for reprocessing",
        "document_id": document_id,
        "chunk_size": settings.CHUNK_SIZE,
        "chunk_overlap": settings.CHUNK_OVERLAP,
        "rechunk": rechunk
    }


class ReparseRequest(BaseModel):
    document_ids: List[int]
    chunk_size: Optional[int] = None
    chunk_overlap: Optional[int] = None
    rechunk: bool = True


@router.post("/reparse-batch")
async def reparse_documents_batch(
    request: ReparseRequest,
    db: AsyncSession = Depends(get_db)
):
    """Batch reparse documents with optional new chunk settings (sequential processing)"""
    # Check batch limit
    if len(request.document_ids) > 20:
        raise HTTPException(status_code=400, detail="Batch limit is 20 documents")

    results = []
    success_count = 0
    failed_count = 0

    # Update chunk settings if provided
    if request.chunk_size is not None:
        settings.CHUNK_SIZE = request.chunk_size
    if request.chunk_overlap is not None:
        settings.CHUNK_OVERLAP = request.chunk_overlap

    for doc_id in request.document_ids:
        try:
            result = await db.execute(
                select(Document).where(Document.id == doc_id)
            )
            document = result.scalar_one_or_none()

            if not document:
                results.append({
                    "document_id": doc_id,
                    "status": "failed",
                    "error": "Document not found"
                })
                failed_count += 1
                continue

            # Reset status for reprocessing
            document.status = "pending"
            document.vector_ids = []
            await db.commit()

            # Process document SYNCHRONOUSLY (not in parallel) to avoid race conditions
            from app.services.document_service import process_document
            success = await process_document(document.id, db, force_rechunk=request.rechunk)
            
            if success:
                results.append({
                    "document_id": doc_id,
                    "document_name": document.name,
                    "status": "completed"
                })
                success_count += 1
            else:
                results.append({
                    "document_id": doc_id,
                    "document_name": document.name,
                    "status": "failed",
                    "error": "Processing failed"
                })
                failed_count += 1

        except Exception as e:
            results.append({
                "document_id": doc_id,
                "status": "failed",
                "error": str(e)
            })
            failed_count += 1

    return {
        "total": len(request.document_ids),
        "success_count": success_count,
        "failed_count": failed_count,
        "results": results
    }


def _overwrite_original_file(file_path: str, file_type: str, content: str) -> str:
    """Overwrite original file with new content. Returns new file path."""
    from docx import Document as DocxDocument
    
    original_path = Path(file_path)
    
    if file_type == "docx":
        doc = DocxDocument()
        for para in content.split('\n'):
            doc.add_paragraph(para)
        doc.save(file_path)
        logger.info(f"Overwritten DOCX file: {file_path}")
        return file_path
    
    elif file_type == "pdf":
        txt_path = original_path.with_suffix('.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Saved edited content as TXT (PDF cannot be overwritten directly): {txt_path}")
        return str(txt_path)
    
    return file_path


class UpdateContentRequest(BaseModel):
    content: str
    chunk_size: Optional[int] = None
    chunk_overlap: Optional[int] = None
    rechunk: bool = True
    overwrite_file: bool = True


@router.put("/{document_id}/content")
async def update_document_content(
    document_id: int,
    request: UpdateContentRequest,
    db: AsyncSession = Depends(get_db)
):
    """Update document content, overwrite original file, and reprocess"""
    logger.info(f"Updating document {document_id} content, overwrite_file={request.overwrite_file}...")
    
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        logger.error(f"Document {document_id} not found")
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        # Update chunk settings if provided
        if request.chunk_size is not None:
            settings.CHUNK_SIZE = request.chunk_size
            logger.info(f"Chunk size updated to {request.chunk_size}")
        if request.chunk_overlap is not None:
            settings.CHUNK_OVERLAP = request.chunk_overlap
            logger.info(f"Chunk overlap updated to {request.chunk_overlap}")

        # Overwrite original file if requested
        file_saved = False
        if request.overwrite_file and document.file_path:
            try:
                new_file_path = _overwrite_original_file(
                    document.file_path,
                    document.file_type,
                    request.content
                )
                document.file_path = new_file_path
                file_saved = True
                logger.info(f"Original file overwritten: {new_file_path}")
            except Exception as e:
                logger.error(f"Failed to overwrite original file: {e}")
                file_saved = False

        # Update document content
        document.content = request.content
        document.status = "pending"
        document.vector_ids = []
        await db.commit()
        logger.info(f"Document {document_id} content updated, status set to pending")

        # Trigger reprocessing
        from app.services.tasks import process_document_async
        import asyncio
        asyncio.create_task(process_document_async(document.id, force_rechunk=request.rechunk))
        logger.info(f"Reprocessing task created for document {document_id}")

        return {
            "message": "Document content updated and queued for reprocessing",
            "document_id": document_id,
            "word_count": len(request.content.split()),
            "file_overwritten": file_saved
        }
    except Exception as e:
        logger.error(f"Error updating document {document_id}: {e}")
        raise


@router.get("/{document_id}/status")
async def get_document_status(
    document_id: int,
    db: AsyncSession = Depends(get_db)
):
    """Get document processing status"""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "document_id": document_id,
        "document_name": document.name,
        "status": document.status,
        "error_message": document.error_message
    }